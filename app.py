import google.generativeai as genai
import streamlit as st
from pathlib import Path
import io
from docx import Document
from PIL import Image
import fitz  # PyMuPDF to convert PDFs to images
from dotenv import load_dotenv
import os

def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

local_css("style.css")
# Load environment variables from .env file
load_dotenv()

# Access API key from environment variable
api_key = os.getenv('GOOGLE_API_KEY')

# Configure GenAI API key
genai.configure(api_key=api_key)

# Function to initialize the model
def initialize_model():
    generation_config = {"temperature": 0.9}
    return genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)

# Function to process the image and generate content based on prompts
def generate_content(model, image_path, prompts, user_prompts):
    image_part = {
        "mime_type": "image/jpeg",
        "data": image_path.getvalue()  # Use getvalue() for BytesIO objects
    }
    
    results = []
    for idx, prompt_text in enumerate(prompts):
        prompt_parts = [prompt_text, image_part]
        response = model.generate_content(prompt_parts)
        
        # Extract and return the text content from the response
        if response.candidates:
            candidate = response.candidates[0]
            if candidate.content and candidate.content.parts:
                text_part = candidate.content.parts[0]
                if text_part.text:
                    # Use the user prompt for display and append description
                    results.append(f"Prompt: {user_prompts[idx]}\nDescription:\n{text_part.text}\n")
                else:
                    results.append(f"Prompt: {user_prompts[idx]}\nDescription: No valid content generated.\n")
            else:
                results.append(f"Prompt: {user_prompts[idx]}\nDescription: No content parts found.\n")
        else:
            results.append(f"Prompt: {user_prompts[idx]}\nDescription: No candidates found.\n")
    
    return results

# Function to create a Word document from the descriptions
def create_word_file(results):
    doc = Document()
    doc.add_heading('Generated Descriptions', 0)

    for description in results:
        doc.add_paragraph(description)

    # Save the Word file to a BytesIO object
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)

    return word_file

# Function to display results with prompt and description
def display_results(results):
    st.write("Medical Insights:")
    for description in results:
        # Separate the prompt and description for better readability
        if "Prompt:" in description:
            prompt_start = description.index("Prompt:")
            description_part = description[prompt_start:].replace("Prompt:", "")
            description_cleaned = description_part.split("Description:", 1)

            # Display the prompt (only the user prompt) on the next line for clarity
            st.write(f"**Prompt**: {description_cleaned[0].strip()}")
            st.write(f"**Description**: {description_cleaned[1].strip()}")
        else:
            st.write(description)

# Function to convert PDF pages to images using PyMuPDF (fitz)
def pdf_to_images(pdf_file):
    images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)  # Get the page
        pix = page.get_pixmap()  # Render the page to an image (Pixmap)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)  # Convert to PIL Image
        
        img_io = io.BytesIO()
        img.save(img_io, format="JPEG")  # Save as JPEG to BytesIO
        img_io.seek(0)
        images.append(img_io)
    
    return images

# Streamlit app
def main():
    # Initialize session state for prompts and results
    if "prompts" not in st.session_state:
        st.session_state.prompts = ""
    if "results" not in st.session_state:
        st.session_state.results = []
    if "uploaded_file" not in st.session_state: 
        st.session_state.uploaded_file = None
    if "history" not in st.session_state:
        st.session_state.history = []

    # Hidden predefined prompt for medical analysis (this will not be shown to the user)
    predefined_prompt = (
        "This system is a medical bot designed for use by doctors only. It uses Vision AI to interpret medical images, "
        "such as X-ray, CT, MRI scans, and medical reports, and provides detailed descriptions and insights using medical terminology. "
        "Please analyze the uploaded image and provide relevant medical insights or descriptions in medical terms. "
        "If the image or input is not related to medical scans, reports, or anything medical, do not respond."
    )

    # Sidebar for navigation
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Choose Functionality", ["Upload Images", "Upload PDF", "History"])

    if page == "Upload Images":
        st.title("Med-Vision: Upload Medical Image (X-ray, CT, MRI)")

        # Upload an image file
        uploaded_file = st.file_uploader("Choose a medical scan (X-ray, CT, MRI)", type=["jpg", "jpeg", "png"])

        if uploaded_file is not None:
            st.session_state.uploaded_file = uploaded_file
            # Save the uploaded image to a BytesIO object
            img_io = io.BytesIO(uploaded_file.getvalue())
            
            # Initialize the model
            model = initialize_model()
            
            # Input for multiple prompts
            st.session_state.prompts = st.text_area(
                "Enter prompts (optional):",
                value=st.session_state.prompts
            )
            
            # Button to generate content
            if st.button("Generate Description"):
                # Split user prompts into a list (these will be displayed)
                user_prompts = [prompt.strip() for prompt in st.session_state.prompts.split('\n') if prompt.strip()]
                
                # Use both the predefined prompt and the user prompt for generation
                prompts = [predefined_prompt] + user_prompts

                st.session_state.results = generate_content(model, img_io, prompts, ["Predefined Medical Prompt"] + user_prompts)

                # Save to history
                st.session_state.history.append({
                    "image": uploaded_file,
                    "results": st.session_state.results
                })

        # Display the uploaded image and previously generated results
        if st.session_state.uploaded_file and st.session_state.results:
            st.image(st.session_state.uploaded_file, caption='Uploaded Medical Scan.', use_column_width=True)
            display_results(st.session_state.results)

            # Create a Word document from the results and provide a download link
            word_file = create_word_file(st.session_state.results)
            st.download_button(
                label="Download Descriptions as Word",
                data=word_file,
                file_name="medical_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif page == "Upload PDF":
        st.title("Med-Vision Advanced -- Upload PDF (e.g., Medical Reports or Scans)")

        # Upload a PDF file
        uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

        if uploaded_file is not None:
            st.session_state.uploaded_file = uploaded_file
            # Convert PDF to images
            images = pdf_to_images(uploaded_file)

            # Initialize the model
            model = initialize_model()

            # Input for multiple prompts
            st.session_state.prompts = st.text_area(
                "Enter prompts (optional):",
                value=st.session_state.prompts
            )

            # Button to generate content
            if st.button("Generate Description"):
                # Split user prompts into a list (these will be displayed)
                user_prompts = [prompt.strip() for prompt in st.session_state.prompts.split('\n') if prompt.strip()]
                
                # Use both the predefined prompt and the user prompt for generation
                prompts = [predefined_prompt] + user_prompts

                all_results = []
                
                # Process all images (from the PDF)
                for img in images:
                    st.session_state.results = generate_content(model, img, prompts, ["Predefined Medical Prompt"] + user_prompts)
                    all_results.extend(st.session_state.results)

                # Save to history
                st.session_state.history.append({
                    "image": uploaded_file,
                    "results": all_results
                })

            # Display images from PDF
            captions = [f'Page {i+1}' for i in range(len(images))]
            st.image(images, caption=captions, use_column_width=True)

        # Display the previously generated results
        if st.session_state.uploaded_file and st.session_state.results:
            display_results(st.session_state.results)

            # Create a Word document from the results and provide a download link
            word_file = create_word_file(st.session_state.results)
            st.download_button(
                label="Download Descriptions as Word",
                data=word_file,
                file_name="medical_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif page == "History":
        st.title("History of Generated Descriptions")
        if st.session_state.history:
            for idx, entry in enumerate(st.session_state.history):
                st.write(f"Entry {idx+1}")
                st.image(entry["image"], caption=f'Image {idx+1}', use_column_width=True)
                for description in entry["results"]:
                    st.write(description)
        else:
            st.write("No history available yet.")

if __name__ == "__main__":
    main()